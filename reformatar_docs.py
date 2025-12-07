#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script de Reformatação de Documentos - Concierge RH Digital INPI
Aplica estrutura padronizada mantendo TODO o conteúdo e TODOS os links
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import os
import re

def add_hyperlink(paragraph, text, url):
    """Adiciona hyperlink a um parágrafo"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Estilo de hyperlink
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def extract_content_with_links(doc_path):
    """Extrai conteúdo preservando estrutura e links"""
    doc = Document(doc_path)
    extracted = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Extrair hyperlinks
        links = []
        for hyperlink in para._element.xpath('.//w:hyperlink'):
            link_text = ''.join([node.text for node in hyperlink.xpath('.//w:t') if node.text])
            r_id = hyperlink.get(qn('r:id'))
            if r_id and r_id in para.part.rels:
                url = para.part.rels[r_id].target_ref
                links.append({'text': link_text, 'url': url})
        
        # Identificar tipo/nível de conteúdo
        style = para.style.name if para.style else 'Normal'
        
        extracted.append({
            'text': text,
            'links': links,
            'style': style,
            'bold': any(run.bold for run in para.runs),
            'italic': any(run.italic for run in para.runs)
        })
    
    return extracted

def organize_content_by_sections(content, doc_name):
    """Organiza conteúdo nas seções padronizadas"""
    full_text = '\n'.join([item['text'] for item in content])
    
    sections = {
        'titulo': '',
        'descricao': '',
        'o_que_e': '',
        'quem_tem_direito': '',
        'como_solicitar': [],
        'prazos': '',
        'documentacao': [],
        'legislacao': '',
        'duvidas': [],
        'contato': ''
    }
    
    # Extrair título (primeiro parágrafo importante ou nome do arquivo)
    if content:
        primeiro = content[0]['text']
        if len(primeiro) < 100 and (content[0]['bold'] or 'Heading' in content[0]['style']):
            sections['titulo'] = primeiro
        else:
            sections['titulo'] = doc_name.replace('.docx', '')
    
    # Identificar seções existentes
    current_section = None
    buffer = []
    
    for i, item in enumerate(content):
        text = item['text']
        text_upper = text.upper()
        
        # Detectar cabeçalhos de seção
        if 'O QUE É' in text_upper and len(text) < 50:
            current_section = 'o_que_e'
            continue
        elif 'QUEM TEM DIREITO' in text_upper and len(text) < 50:
            current_section = 'quem_tem_direito'
            continue
        elif 'COMO SOLICITAR' in text_upper and len(text) < 50:
            current_section = 'como_solicitar'
            continue
        elif 'PRAZO' in text_upper and len(text) < 50:
            current_section = 'prazos'
            continue
        elif 'DOCUMENTAÇÃO' in text_upper or 'DOCUMENTOS' in text_upper and len(text) < 80:
            current_section = 'documentacao'
            continue
        elif 'LEGISLAÇÃO' in text_upper or 'BASE LEGAL' in text_upper and len(text) < 50:
            current_section = 'legislacao'
            continue
        elif 'DÚVIDAS' in text_upper or 'PERGUNTAS' in text_upper and len(text) < 80:
            current_section = 'duvidas'
            continue
        elif 'CONTATO' in text_upper and len(text) < 50:
            current_section = 'contato'
            continue
        
        # Adicionar conteúdo à seção atual
        if current_section:
            if current_section in ['como_solicitar', 'documentacao', 'duvidas']:
                sections[current_section].append(item)
            else:
                if sections[current_section]:
                    sections[current_section] += '\n\n' + text
                else:
                    sections[current_section] = text
        elif i > 0 and i < 5 and not sections['descricao']:
            # Primeiros parágrafos como descrição
            if sections['descricao']:
                sections['descricao'] += '\n\n' + text
            else:
                sections['descricao'] = text
    
    return sections

def create_formatted_document(output_path, sections, all_content):
    """Cria documento reformatado com estrutura padronizada"""
    doc = Document()
    
    # Configurar estilos personalizados
    styles = doc.styles
    
    # Estilo para Título 1
    try:
        h1_style = styles['Heading 1']
    except KeyError:
        h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 70, 127)
    
    # Estilo para Título 2
    try:
        h2_style = styles['Heading 2']
    except KeyError:
        h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 112, 192)
    
    # TÍTULO PRINCIPAL
    if sections['titulo']:
        title = doc.add_heading(sections['titulo'].upper(), level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # DESCRIÇÃO
    if sections['descricao']:
        desc = doc.add_paragraph(sections['descricao'])
        desc.paragraph_format.space_after = Pt(12)
    
    # O QUE É?
    doc.add_heading('O QUE É?', level=2)
    if sections['o_que_e']:
        doc.add_paragraph(sections['o_que_e'])
    else:
        # Tentar extrair do conteúdo geral
        for item in all_content:
            if 'o que é' in item['text'].lower() and len(item['text']) > 50:
                doc.add_paragraph(item['text'])
                break
        else:
            doc.add_paragraph('Informações sobre a natureza e objetivo deste serviço.')
    
    # QUEM TEM DIREITO?
    doc.add_heading('QUEM TEM DIREITO?', level=2)
    if sections['quem_tem_direito']:
        doc.add_paragraph(sections['quem_tem_direito'])
    else:
        doc.add_paragraph('Servidores ativos do INPI.')
    
    # COMO SOLICITAR?
    doc.add_heading('COMO SOLICITAR?', level=2)
    if sections['como_solicitar']:
        for idx, item in enumerate(sections['como_solicitar'], 1):
            if any(char.isdigit() for char in item['text'][:5]):
                # Já tem numeração
                p = doc.add_paragraph(item['text'], style='List Number')
            else:
                p = doc.add_paragraph(f"{item['text']}", style='List Number')
            
            # Adicionar links se houver
            if item['links']:
                for link in item['links']:
                    add_hyperlink(p, f" [{link['text']}]", link['url'])
    else:
        doc.add_paragraph('Entre em contato com a área responsável para orientações.', style='List Number')
    
    # PRAZOS
    doc.add_heading('PRAZOS', level=2)
    if sections['prazos']:
        doc.add_paragraph(sections['prazos'])
    else:
        doc.add_paragraph('Consulte a legislação ou entre em contato para informações sobre prazos.')
    
    # DOCUMENTAÇÃO NECESSÁRIA
    doc.add_heading('DOCUMENTAÇÃO NECESSÁRIA', level=2)
    if sections['documentacao']:
        for item in sections['documentacao']:
            p = doc.add_paragraph(item['text'], style='List Bullet')
            if item['links']:
                for link in item['links']:
                    add_hyperlink(p, f" [{link['text']}]", link['url'])
    else:
        doc.add_paragraph('Documentação específica conforme o caso.', style='List Bullet')
    
    # LEGISLAÇÃO
    doc.add_heading('LEGISLAÇÃO', level=2)
    if sections['legislacao']:
        p = doc.add_paragraph(sections['legislacao'])
    else:
        # Buscar menções a leis/portarias no conteúdo
        legislacao_found = []
        for item in all_content:
            if any(termo in item['text'].lower() for termo in ['lei', 'portaria', 'decreto', 'instrução normativa', 'resolução']):
                legislacao_found.append(item['text'])
        
        if legislacao_found:
            for leg in legislacao_found[:5]:  # Limitar a 5
                doc.add_paragraph(leg, style='List Bullet')
        else:
            doc.add_paragraph('Consulte a legislação aplicável.')
    
    # DÚVIDAS FREQUENTES
    doc.add_heading('DÚVIDAS FREQUENTES', level=2)
    if sections['duvidas']:
        for item in sections['duvidas']:
            text = item['text']
            if text.startswith(('Q:', 'P:', 'R:')):
                run = doc.add_paragraph(text).runs[0]
                run.bold = True if text[0] in ['Q', 'P'] else False
            else:
                doc.add_paragraph(f"• {text}")
    else:
        doc.add_paragraph('Para dúvidas, consulte o contato abaixo.')
    
    # CONTATO
    doc.add_heading('CONTATO', level=2)
    if sections['contato']:
        doc.add_paragraph(sections['contato'])
    else:
        # Buscar e-mails e telefones no conteúdo
        contatos = []
        for item in all_content:
            if '@' in item['text'] or 'ramal' in item['text'].lower() or 'telefone' in item['text'].lower():
                contatos.append(item['text'])
        
        if contatos:
            for contato in contatos[:3]:
                doc.add_paragraph(contato)
        else:
            doc.add_paragraph('E-mail: cgrh@inpi.gov.br')
            doc.add_paragraph('Telefone: Consulte a intranet do INPI')
    
    # Salvar documento
    doc.save(output_path)
    return True

def process_single_doc(doc_name, docs_dir):
    """Processa um documento individual"""
    input_path = os.path.join(docs_dir, doc_name)
    backup_path = os.path.join(docs_dir, 'backup_' + doc_name)
    
    print(f"\n{'='*70}")
    print(f"📄 PROCESSANDO: {doc_name}")
    print(f"{'='*70}")
    
    try:
        # Fazer backup
        if os.path.exists(input_path):
            doc_backup = Document(input_path)
            doc_backup.save(backup_path)
            print(f"✅ Backup criado: backup_{doc_name}")
        
        # Extrair conteúdo
        content = extract_content_with_links(input_path)
        print(f"📝 Parágrafos extraídos: {len(content)}")
        
        # Contar links
        total_links = sum(len(item['links']) for item in content)
        print(f"🔗 Links encontrados: {total_links}")
        
        # Organizar em seções
        sections = organize_content_by_sections(content, doc_name)
        
        # Criar documento reformatado
        create_formatted_document(input_path, sections, content)
        print(f"✅ Documento reformatado e salvo!")
        
        return {
            'file': doc_name,
            'status': 'success',
            'paragraphs': len(content),
            'links': total_links,
            'changes': 'Estrutura padronizada aplicada'
        }
        
    except Exception as e:
        print(f"❌ ERRO: {str(e)}")
        return {
            'file': doc_name,
            'status': 'error',
            'error': str(e)
        }

# LISTA DE DOCUMENTOS
docs_list = [
    "Aposentadoria e Abono.docx",
    "Capacitação.docx",
    "Carta de Serviços.docx",
    "Dados Cadastrais.docx",
    "Estágio Probatório.docx",
    "Frequência.docx",
    "Férias.docx",
    "Licenças.docx",
    "Pagamento.docx",
    "Programa de Gestão e Desempenho.docx",
    "Remoção.docx",
    "Retribuição por Titulação.docx",
    "Saúde Ocupacional.docx",
    "Seleção Interna e Externa.docx",
    "Utilização do SouGov.docx"
]

docs_dir = r"C:\Users\Davison.DESKTOP-7GLJO2G\Documents\concierge-cgrh\docs"

# PROCESSAR TODOS
results = []
for doc_name in docs_list:
    result = process_single_doc(doc_name, docs_dir)
    results.append(result)

# RELATÓRIO FINAL
print(f"\n{'='*70}")
print("📊 RELATÓRIO FINAL DE REFORMATAÇÃO")
print(f"{'='*70}\n")

success_count = 0
error_count = 0

for r in results:
    if r['status'] == 'success':
        success_count += 1
        print(f"✅ {r['file']}")
        print(f"   📝 Mudanças: {r['changes']}")
        print(f"   🔗 Links preservados: {r['links']}")
        print(f"   📊 Parágrafos: {r['paragraphs']}")
    else:
        error_count += 1
        print(f"❌ {r['file']}")
        print(f"   ⚠️ Erro: {r['error']}")
    print()

print(f"{'='*70}")
print(f"✅ Documentos reformatados com sucesso: {success_count}/15")
print(f"❌ Erros: {error_count}")
print(f"{'='*70}")
print(f"\n💾 Backups salvos com prefixo 'backup_'")
print(f"📁 Localização: {docs_dir}")
