#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script FINAL de Reformatação - Preservação TOTAL de Links
Corrige a preservação de hyperlinks durante a reformatação
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import os
import re

def add_hyperlink(paragraph, url, text=None):
    """Adiciona hyperlink funcional a um parágrafo"""
    # Se text não especificado, usa a URL
    if text is None:
        text = url
    
    # Obter a parte do documento
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Criar elemento hyperlink
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Criar run para o texto
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Estilo azul sublinhado
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    # Adicionar texto
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def extract_all_hyperlinks(doc):
    """Extrai TODOS os hyperlinks do documento"""
    all_links = []
    
    for para in doc.paragraphs:
        for hyperlink in para._element.xpath('.//w:hyperlink'):
            link_text = ''.join([node.text for node in hyperlink.xpath('.//w:t') if node.text])
            r_id = hyperlink.get(qn('r:id'))
            
            if r_id and r_id in para.part.rels:
                url = para.part.rels[r_id].target_ref
                all_links.append({
                    'text': link_text,
                    'url': url,
                    'context': para.text[:100]
                })
    
    # Extrair de tabelas também
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for hyperlink in para._element.xpath('.//w:hyperlink'):
                        link_text = ''.join([node.text for node in hyperlink.xpath('.//w:t') if node.text])
                        r_id = hyperlink.get(qn('r:id'))
                        
                        if r_id and r_id in para.part.rels:
                            url = para.part.rels[r_id].target_ref
                            all_links.append({
                                'text': link_text,
                                'url': url,
                                'context': para.text[:100]
                            })
    
    return all_links

def restore_links_to_document(doc_path):
    """Restaura links do backup para o documento reformatado"""
    backup_path = doc_path.replace('.docx', '_BACKUP_ORIGINAL.docx')
    
    # Se não existe backup com este nome, tentar o outro formato
    if not os.path.exists(backup_path):
        backup_path = os.path.join(
            os.path.dirname(doc_path),
            'backup_' + os.path.basename(doc_path)
        )
    
    if not os.path.exists(backup_path):
        print(f"   ⚠️ Backup não encontrado para restaurar links")
        return 0
    
    try:
        # Extrair links do backup
        backup_doc = Document(backup_path)
        original_links = extract_all_hyperlinks(backup_doc)
        
        if not original_links:
            return 0
        
        print(f"   🔗 Links no original: {len(original_links)}")
        
        # Abrir documento reformatado
        doc = Document(doc_path)
        links_added = 0
        
        # Para cada link original, tentar adicionar ao documento reformatado
        for link_info in original_links:
            url = link_info['url']
            text = link_info['text']
            context = link_info['context']
            
            # Procurar parágrafo similar no documento reformatado
            for para in doc.paragraphs:
                para_text = para.text
                
                # Se encontrar contexto similar e o link ainda não está lá
                if (text in para_text or context[:50] in para_text):
                    # Verificar se já tem este link
                    existing = False
                    for existing_link in para._element.xpath('.//w:hyperlink'):
                        r_id = existing_link.get(qn('r:id'))
                        if r_id and r_id in para.part.rels:
                            if para.part.rels[r_id].target_ref == url:
                                existing = True
                                break
                    
                    if not existing:
                        # Adicionar link ao final do parágrafo
                        para.add_run(' ')
                        add_hyperlink(para, url, text)
                        links_added += 1
                        break
        
        # Salvar documento com links restaurados
        doc.save(doc_path)
        print(f"   ✅ Links restaurados: {links_added}")
        
        return links_added
        
    except Exception as e:
        print(f"   ❌ Erro ao restaurar links: {str(e)}")
        return 0

def compare_links(doc_name, docs_dir):
    """Compara links entre original e reformatado"""
    backup_path = os.path.join(docs_dir, 'backup_' + doc_name)
    current_path = os.path.join(docs_dir, doc_name)
    
    if not os.path.exists(backup_path):
        return None
    
    try:
        # Links no backup
        backup_doc = Document(backup_path)
        backup_links = extract_all_hyperlinks(backup_doc)
        
        # Links no reformatado
        current_doc = Document(current_path)
        current_links = extract_all_hyperlinks(current_doc)
        
        return {
            'original': len(backup_links),
            'current': len(current_links),
            'missing': len(backup_links) - len(current_links),
            'original_links': backup_links,
            'current_links': current_links
        }
    except Exception as e:
        return {'error': str(e)}

# Lista de documentos
docs_list = [
    "Aposentadoria e Abono.docx",
    "Capacitação.docx",
    "Carta de Serviços.docx",
    "Dados Cadastrais.docx",
    "Estágio Probatório.docx",
    "Frequência.docx",
    "Férias.docx",
    "Licenças.docx",
    "Pagamento.docx",
    "Programa de Gestão e Desempenho.docx",
    "Remoção.docx",
    "Retribuição por Titulação.docx",
    "Saúde Ocupacional.docx",
    "Seleção Interna e Externa.docx",
    "Utilização do SouGov.docx"
]

docs_dir = r"C:\Users\Davison.DESKTOP-7GLJO2G\Documents\concierge-cgrh\docs"

print("="*80)
print("ANÁLISE DE LINKS - ORIGINAL vs REFORMATADO")
print("="*80)
print()

docs_with_missing_links = []

for doc_name in docs_list:
    print(f"📄 {doc_name}")
    
    comparison = compare_links(doc_name, docs_dir)
    
    if comparison and 'error' not in comparison:
        print(f"   Original: {comparison['original']} links")
        print(f"   Reformatado: {comparison['current']} links")
        
        if comparison['missing'] > 0:
            print(f"   ⚠️ FALTAM {comparison['missing']} links!")
            docs_with_missing_links.append({
                'name': doc_name,
                'missing': comparison['missing'],
                'original_links': comparison['original_links']
            })
            
            # Mostrar links que faltam
            print(f"   📋 Links originais:")
            for link in comparison['original_links']:
                print(f"      • {link['text'][:50]} -> {link['url'][:60]}")
        else:
            print(f"   ✅ Todos os links preservados!")
    
    print()

if docs_with_missing_links:
    print("="*80)
    print("RESTAURANDO LINKS FALTANTES")
    print("="*80)
    print()
    
    for doc_info in docs_with_missing_links:
        doc_name = doc_info['name']
        doc_path = os.path.join(docs_dir, doc_name)
        
        print(f"🔧 Restaurando: {doc_name}")
        restored = restore_links_to_document(doc_path)
        print()

print("="*80)
print("✅ ANÁLISE E RESTAURAÇÃO CONCLUÍDAS")
print("="*80)
