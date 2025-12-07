#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script para reformatar documentos Word do Concierge RH Digital
Mantém TODO o conteúdo original e preserva TODOS os links
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import os
import re

def extract_hyperlinks(paragraph):
    """Extrai todos os hyperlinks de um parágrafo"""
    links = []
    for rel in paragraph.part.rels.values():
        if "hyperlink" in rel.target_ref:
            links.append(rel.target_ref)
    
    # Extrair hyperlinks do XML
    hyperlinks_data = []
    for hyperlink in paragraph._element.xpath('.//w:hyperlink'):
        text_content = ''.join([node.text for node in hyperlink.xpath('.//w:t')])
        r_id = hyperlink.get(qn('r:id'))
        if r_id and r_id in paragraph.part.rels:
            url = paragraph.part.rels[r_id].target_ref
            hyperlinks_data.append({'text': text_content, 'url': url})
    
    return hyperlinks_data

def extract_all_content(doc_path):
    """Extrai todo o conteúdo do documento preservando links"""
    try:
        doc = Document(doc_path)
        content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Extrair hyperlinks
                links = extract_hyperlinks(para)
                content.append({
                    'text': text,
                    'links': links,
                    'style': para.style.name
                })
        
        # Extrair tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            links = extract_hyperlinks(para)
                            content.append({
                                'text': text,
                                'links': links,
                                'style': 'Table'
                            })
        
        return content
    except Exception as e:
        print(f"Erro ao ler {doc_path}: {str(e)}")
        return []

def create_formatted_doc(output_path, title, sections):
    """Cria documento formatado com estrutura padronizada"""
    doc = Document()
    
    # Configurar estilos
    styles = doc.styles
    
    # Título principal
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Adicionar seções
    for section_title, section_content in sections.items():
        if section_content:
            # Adicionar título da seção
            doc.add_heading(section_title, level=2)
            
            # Adicionar conteúdo
            if isinstance(section_content, list):
                for item in section_content:
                    p = doc.add_paragraph(item, style='List Bullet')
            else:
                p = doc.add_paragraph(section_content)
    
    # Salvar documento
    doc.save(output_path)
    return True

def analyze_document_structure(content):
    """Analisa conteúdo e identifica seções"""
    full_text = '\n'.join([item['text'] for item in content])
    
    sections = {
        'Descrição': '',
        'O QUE É?': '',
        'QUEM TEM DIREITO?': '',
        'COMO SOLICITAR?': '',
        'PRAZOS': '',
        'DOCUMENTAÇÃO NECESSÁRIA': '',
        'LEGISLAÇÃO': '',
        'DÚVIDAS FREQUENTES': '',
        'CONTATO': ''
    }
    
    return sections, full_text

def process_document(input_path, output_path):
    """Processa um documento individual"""
    print(f"\n{'='*60}")
    print(f"Processando: {os.path.basename(input_path)}")
    print(f"{'='*60}")
    
    # Extrair conteúdo
    content = extract_all_content(input_path)
    
    if not content:
        print("⚠️ Nenhum conteúdo extraído!")
        return None
    
    # Contar links
    total_links = sum(len(item['links']) for item in content)
    
    # Extrair texto completo
    full_text = '\n\n'.join([item['text'] for item in content])
    
    print(f"📄 Conteúdo extraído: {len(content)} parágrafos")
    print(f"🔗 Links encontrados: {total_links}")
    print(f"\n--- PREVIEW DO CONTEÚDO ---")
    print(full_text[:500] + "..." if len(full_text) > 500 else full_text)
    
    return {
        'content': content,
        'links': total_links,
        'text': full_text
    }

# Lista de documentos
docs = [
    "Aposentadoria e Abono.docx",
    "Capacitação.docx",
    "Carta de Serviços.docx",
    "Dados Cadastrais.docx",
    "Estágio Probatório.docx",
    "Frequência.docx",
    "Férias.docx",
    "Licenças.docx",
    "Pagamento.docx",
    "Programa de Gestão e Desempenho.docx",
    "Remoção.docx",
    "Retribuição por Titulação.docx",
    "Saúde Ocupacional.docx",
    "Seleção Interna e Externa.docx",
    "Utilização do SouGov.docx"
]

# Diretório dos documentos
docs_dir = r"C:\Users\Davison.DESKTOP-7GLJO2G\Documents\concierge-cgrh\docs"

# Processar todos os documentos
results = []
for doc_name in docs:
    input_path = os.path.join(docs_dir, doc_name)
    if os.path.exists(input_path):
        result = process_document(input_path, input_path)
        if result:
            results.append({
                'file': doc_name,
                'links': result['links'],
                'paragraphs': len(result['content'])
            })
    else:
        print(f"❌ Arquivo não encontrado: {doc_name}")

# Relatório final
print(f"\n{'='*60}")
print("RELATÓRIO FINAL - ANÁLISE DE CONTEÚDO")
print(f"{'='*60}")
for r in results:
    print(f"✅ {r['file']}")
    print(f"   📝 Parágrafos: {r['paragraphs']}")
    print(f"   🔗 Links: {r['links']}")
print(f"{'='*60}")
print(f"Total de documentos analisados: {len(results)}/15")
